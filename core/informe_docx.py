"""
Generación del informe individual de un alumno en Word (.docx), a partir
de las estructuras recopiladas en core/informe_alumno.py.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from core.informe_alumno import InformeAlumno, InformeAlumnoFinal

COLOR_CABECERA = RGBColor(0x0F, 0x6F, 0xB9)
COLOR_RESULTADO = RGBColor(0x0D, 0x3D, 0x6B)

_ETIQUETAS_EVALUACION = {
    "1EVA": "Primera evaluación",
    "2EVA": "Segunda evaluación",
    "3EVA": "Tercera evaluación",
    "FINAL": "Evaluación final de curso",
}


def _formatear_numero(valor: float | None) -> str:
    if valor is None:
        return "—"
    return f"{valor:.2f}".replace(".", ",")


def _sombrear_celda(celda, color_hex: str):
    """Aplica un color de fondo a una celda de tabla (python-docx no tiene
    un método directo para esto; se manipula el XML interno de la celda).
    """
    propiedades_celda = celda._tc.get_or_add_tcPr()
    sombreado = propiedades_celda.makeelement(qn("w:shd"), {qn("w:fill"): color_hex})
    propiedades_celda.append(sombreado)


def _anadir_tabla(documento: Document, encabezados: list[str], filas: list[list[str]]):
    tabla = documento.add_table(rows=1, cols=len(encabezados))
    tabla.style = "Table Grid"
    fila_cabecera = tabla.rows[0]
    for celda, texto in zip(fila_cabecera.cells, encabezados):
        celda.text = texto
        _sombrear_celda(celda, "0F6FB9")
        for parrafo in celda.paragraphs:
            parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in parrafo.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for indice_fila, datos_fila in enumerate(filas):
        fila_tabla = tabla.add_row()
        if indice_fila % 2 == 1:
            for celda in fila_tabla.cells:
                _sombrear_celda(celda, "EAF6F4")
        for celda, valor in zip(fila_tabla.cells, datos_fila):
            celda.text = str(valor)
            for parrafo in celda.paragraphs:
                parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return tabla


def _anadir_cabecera_comun(
    documento: Document, nombre_materia: str, etiqueta_evaluacion: str, apellidos: str, nombre: str
):
    titulo = documento.add_heading("Informe de calificaciones", level=0)
    for run in titulo.runs:
        run.font.color.rgb = COLOR_RESULTADO

    parrafo_alumno = documento.add_paragraph()
    parrafo_alumno.add_run("Alumno/a: ").bold = True
    parrafo_alumno.add_run(f"{apellidos}, {nombre}")

    parrafo_materia = documento.add_paragraph()
    parrafo_materia.add_run("Materia: ").bold = True
    parrafo_materia.add_run(nombre_materia)

    parrafo_evaluacion = documento.add_paragraph()
    parrafo_evaluacion.add_run("Evaluación: ").bold = True
    parrafo_evaluacion.add_run(etiqueta_evaluacion)

    documento.add_paragraph()


def _anadir_nota_final(documento: Document, etiqueta: str, valor: float | None, calificacion: str):
    parrafo = documento.add_paragraph()
    run = parrafo.add_run(f"{etiqueta}: {_formatear_numero(valor)} ({calificacion or '—'})")
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = COLOR_RESULTADO
    documento.add_paragraph()


def generar_informe_evaluacion_docx(informe: InformeAlumno, ruta_destino: str | Path) -> Path:
    ruta_destino = Path(ruta_destino)
    documento = Document()

    etiqueta_evaluacion = _ETIQUETAS_EVALUACION.get(informe.nombre_evaluacion, informe.nombre_evaluacion)
    _anadir_cabecera_comun(
        documento, informe.nombre_materia, etiqueta_evaluacion, informe.apellidos_alumno, informe.nombre_alumno
    )
    _anadir_nota_final(
        documento, "Nota final de la evaluación", informe.nota_final_numerica, informe.calificacion_final
    )

    documento.add_heading("Calificación de cada criterio de evaluación", level=1)
    filas = [
        [fila.codigo_criterio, f"{fila.peso_en_instrumento:g}", _formatear_numero(fila.valor), fila.calificacion or "—"]
        for fila in informe.filas_criterios
    ]
    _anadir_tabla(documento, ["Criterio", "Peso en la materia", "Nota", "Calificación"], filas)

    documento.add_paragraph()
    documento.add_heading("Desglose por instrumento de evaluación", level=1)
    if not informe.bloques_instrumentos:
        documento.add_paragraph("No hay instrumentos de evaluación con criterios asignados.")
    for bloque in informe.bloques_instrumentos:
        parrafo_bloque = documento.add_paragraph()
        parrafo_bloque.add_run(f"{bloque.nombre_instrumento}").bold = True
        parrafo_bloque.add_run(f" — peso {bloque.peso_global:g}% de la evaluación")
        filas_bloque = [
            [
                fila.codigo_criterio,
                f"{fila.peso_en_instrumento:g}%",
                _formatear_numero(fila.valor),
                fila.calificacion or "—",
            ]
            for fila in bloque.criterios
        ]
        _anadir_tabla(documento, ["Criterio", "Peso en este instrumento", "Nota", "Calificación"], filas_bloque)
        documento.add_paragraph()

    documento.save(str(ruta_destino))
    return ruta_destino


def generar_informe_final_docx(informe: InformeAlumnoFinal, ruta_destino: str | Path) -> Path:
    ruta_destino = Path(ruta_destino)
    documento = Document()

    _anadir_cabecera_comun(
        documento, informe.nombre_materia, _ETIQUETAS_EVALUACION["FINAL"],
        informe.apellidos_alumno, informe.nombre_alumno,
    )
    _anadir_nota_final(documento, "Nota final de curso", informe.nota_final_numerica, informe.calificacion_final)

    pesos_texto = "  ·  ".join(f"{nombre}: {peso:g}" for nombre, peso in informe.pesos_evaluaciones.items())
    parrafo_pesos = documento.add_paragraph()
    run_pesos = parrafo_pesos.add_run(f"Pesos aplicados entre evaluaciones — {pesos_texto}")
    run_pesos.italic = True
    documento.add_paragraph()

    documento.add_heading("Calificación de cada criterio, por evaluación", level=1)
    filas = [
        [
            fila.codigo_criterio,
            _formatear_numero(fila.valor_1eva),
            _formatear_numero(fila.valor_2eva),
            _formatear_numero(fila.valor_3eva),
            _formatear_numero(fila.valor),
            fila.calificacion or "—",
        ]
        for fila in informe.filas_criterios
    ]
    _anadir_tabla(documento, ["Criterio", "1ª Ev.", "2ª Ev.", "3ª Ev.", "Nota FINAL", "Calif."], filas)

    documento.add_paragraph()
    parrafo_nota = documento.add_paragraph(
        "Las celdas en blanco (—) indican que ese criterio no tuvo ninguna nota calculada en esa "
        "evaluación. La nota FINAL combina las evaluaciones que sí tienen nota, según los pesos "
        "indicados arriba."
    )
    parrafo_nota.runs[0].italic = True

    documento.save(str(ruta_destino))
    return ruta_destino
